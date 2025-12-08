import requests
import json
from docx import Document
import os
import getpass
from dotenv import load_dotenv
from langsmith import traceable
import re

# =============================================================================
#                           SETUP
# =============================================================================

load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    os.environ["OPENAI_API_KEY"] = getpass.getpass("Enter your OpenAI API key: ")

os.environ["LANGSMITH_TRACING"] = "true"
os.environ["LANGSMITH_PROJECT"] = "8D-LLM-Monitoring"

LITELLM_BASE = "http://litellm.ame.local"
MODEL = "azure/gpt-4.1"


# =============================================================================
#                           LLM CALL WRAPPER
# =============================================================================

@traceable(run_type="llm", name="8D-Section-Extraction")
def call_litellm(prompt, model=MODEL, api_base=LITELLM_BASE):
    url = f"{api_base}/v1/chat/completions"

    headers = {
        "Authorization": f"Bearer {os.environ['OPENAI_API_KEY']}",
        "Content-Type": "application/json",
    }

    payload = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0,
        "stream": False,
    }

    resp = requests.post(url, json=payload, headers=headers)
    resp.raise_for_status()
    data = resp.json()
    return data["choices"][0]["message"]["content"]


def safe_json_loads(text):
    """Cleans LLM output and extracts valid JSON."""
    text = text.strip()

    # Remove fenced code
    text = re.sub(r"^```json\s*", "", text)
    text = re.sub(r"^```\s*", "", text)
    text = re.sub(r"\s*```$", "", text)

    # Extract first {...}
    if "{" in text and "}" in text:
        text = text[text.find("{"): text.rfind("}") + 1]

    return json.loads(text)


# =============================================================================
#                           DOC PARSING
# =============================================================================
def split_by_headings(doc_path):
    doc = Document(doc_path)

    sections = []
    current_title = None
    current_text = []

    for para in doc.paragraphs:
        style = para.style.name

        # ------------------------------------------------------------------
        # Heading 2 → Start of a new main section (e.g., "2.2 D2", "2.4 D4")
        # ------------------------------------------------------------------
        if style == "Heading 2":
            # Save previous section
            if current_title:
                sections.append({
                    "title": current_title,
                    "content": "\n".join(current_text).strip()
                })

            current_title = para.text.strip()
            current_text = []  # Reset text collector

        # ------------------------------------------------------------------
        # Heading 3 → belongs to the CURRENT Section!
        # ------------------------------------------------------------------
        elif style == "Heading 3":
            if current_title:
                current_text.append(para.text.strip())

        else:
            # Normal paragraph belongs to the current active section
            if current_title:
                if para.text.strip():
                    current_text.append(para.text.strip())

    # Save last section
    if current_title:
        sections.append({
            "title": current_title,
            "content": "\n".join(current_text).strip()
        })

    return sections



def extract_product_from_tables(doc_path): # For the product name extraction
    doc = Document(doc_path)
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue

            key = row.cells[0].text.strip().lower()
            value = row.cells[1].text.strip()

            if key.startswith("product"):
                return value

    return None



def filter_sections(sections, want_prefixes):
    filtered_section = [
        sec for sec in sections
        if any(sec["title"].startswith(p) for p in want_prefixes)
    ]
    return filtered_section
# =============================================================================
#                           SECTION EXTRACTION
# =============================================================================

# def extract_section_summary(title, content):
#     """General extractor for D1-D8 sections."""

#     prompt = f"""
# You are converting an 8D section into JSON.

# SECTION TITLE:
# {title}

# TEXT:
# {content}

# Extract a concise engineering summary. Return ONLY JSON:

# {{
#   "summary": "<1-3 sentence summary>",
#   "keywords": ["...","..."]
# }}
# """
#     return safe_json_loads(call_litellm(prompt))


def extract_problem_from_D2(content):
    prompt = f"""
You are a FMEA expert, now you read the D2 section of 8D report, extract the system name and corrseponding elements, 
System is a more general term, while element is a specific part of the system like the power supply, PCBA board, etc.
The failure
TEXT:
{content}

Return ONLY:

{{
    "System_name": "",
    "System_element":"",
  "problem_symptoms": "",
  "derived_failure_mode": ""
}}
"""
    return safe_json_loads(call_litellm(prompt))


def extract_root_cause_from_D4(content):
    prompt = f"""
Extract core root cause from D4.

TEXT:
{content}

Return ONLY:

{{
  "root_cause": "",
  "root_cause_chain": ""
}}
"""
    return safe_json_loads(call_litellm(prompt))


# =============================================================================
#                 INFER FMEA HEADER FROM ALL SECTIONS
# =============================================================================

# def infer_fmea_header(all_sections_dict):
#     prompt = f"""
# You are summarizing an entire 8D report into a structured FMEA-style header.

# SECTIONS:
# {json.dumps(all_sections_dict, indent=2)}

# Return ONLY JSON:

# {{
#   "system": "",
#   "system_element": "",
#   "function": "",
#   "failure_mode": "",
#   "failure_cause": "",
#   "failure_effect": ""
# }}
# """
#     return safe_json_loads(call_litellm(prompt))


# =============================================================================
#                           MAIN BUILDER
# =============================================================================

def build_8d_json(doc_path):
    # Extract ID from filename
    doc_name = os.path.basename(doc_path)
    eight_d_id = re.findall(r"8D\d+", doc_name)
    eight_d_id = eight_d_id[0] if eight_d_id else doc_name

    product_name = extract_product_from_tables(doc_path)
    sections_raw = split_by_headings(doc_path)

    # ---- NEW: filter sections ----
    want_prefixes = ["Introduction", "D2",  "D4", ]
    filtered_sections = filter_sections(sections_raw, want_prefixes)

    # optional: debug print
    print("Filtered sections:")
    for sec in filtered_sections:
        print(" -", sec["title"])

    result = {
        "8d_id": eight_d_id,
        "product_name": product_name,
        "sections": {},
        "raw_context": {}
    }

    # ---------------------------------------------------------------------
    # Only D2–D6 go into structured sections
    # ---------------------------------------------------------------------
    for sec in filtered_sections:
        title = sec["title"]
        content = sec["content"]

        # skip Introduction for now (you can later parse system info here)
        if title.startswith("Introduction"):
            continue

        # D2
        if title.startswith("D2"):
            result["raw_context"]["D2"] = content
            result["sections"]["D2"] = extract_problem_from_D2(content)

        # # D3
        # elif title.startswith("D3"):
        #     result["raw_context"]["D3"] = content
        #     result["sections"]["D3"] = extract_section_summary(title, content)

        # D4
        elif title.startswith("D4"):
            result["raw_context"]["D4"] = content
            result["sections"]["D4"] = extract_root_cause_from_D4(content)

        # # D5
        # elif title.startswith("D5"):
        #     result["raw_context"]["D5"] = content
        #     result["sections"]["D5"] = extract_section_summary(title, content)

        # D6
        # elif title.startswith("D6"):
        #     result["raw_context"]["D6"] = content
        #     result["sections"]["D6"] = extract_section_summary(title, content)

    # ---------------------------------------------------------------------
    # FMEA header from extracted sections (D2–D6)
    # ---------------------------------------------------------------------
    # fmea_header = infer_fmea_header(result["sections"])
    # result.update(fmea_header)

    return result


# =============================================================================
#                               RUN
# =============================================================================

if __name__ == "__main__":
    doc_path = r"C:\Users\FW\Desktop\FMEA_AI\Project_Phase\DATA\8D\8D6264240043R03.docx"  
    schema = build_8d_json(doc_path)

    # use DOCX filename as JSON output name
    base_name = os.path.basename(doc_path)          
    stem, _ = os.path.splitext(base_name)          
    output_path = f"{stem}.json"                   

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(schema, f, indent=2, ensure_ascii=False)

    print(f"8D JSON schema generated: {output_path}")