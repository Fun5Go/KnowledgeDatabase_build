from docx import Document
import docx2txt
import re 
import mammoth


# def extract_product_name_mammoth(doc_path):
#     # 用 mammoth 提取所有正文文本，包括文本框、shape、表格、段落
#     with open(doc_path, "rb") as f:
#         text = mammoth.extract_raw_text(f).value

#     print("======= RAW TEXT FROM MAMMOTH =======")
#     print(text)
#     print("=====================================")

#     # 匹配 Product 行
#     for line in text.split("\n"):
#         clean = line.strip()
#         if clean.lower().startswith("product"):
#             parts = clean.split(":", 1)
#             if len(parts) > 1:
#                 return parts[1].strip()

#     return None

def extract_product_from_tables(doc_path):
    doc = Document(doc_path)

    for table in doc.tables:
        for row in table.rows:
            # 必须 >= 2 列
            if len(row.cells) < 2:
                continue

            key = row.cells[0].text.strip().lower()
            value = row.cells[1].text.strip()

            # 找到 Product
            if key.startswith("product"):
                return value

    return None


# if __name__ == "__main__":
doc_path = r"C:\Users\FW\Desktop\FMEA_AI\Project_Phase\DATA\8D\8D6557080300R02.docx" # 改成你的文件路径

product = extract_product_from_tables(doc_path)
print("Product:", product)
#     product = extract_product_name_mammoth(doc_path)
#     print("Product name from Mammoth:", product)
# doc = Document(doc_path)
# print("\n==========================")
# print(" 🔍 TEST 1: paragraphs")
# print("==========================")
# for i, p in enumerate(doc.paragraphs):
#     print(f"[P{i}] '{p.text}'")

# print("\n==========================")
# print(" 🔍 TEST 2: tables")
# print("==========================")
# for ti, table in enumerate(doc.tables):
#     print(f"\n-- Table {ti} --")
#     for ri, row in enumerate(table.rows):
#         row_text = [cell.text for cell in row.cells]
#         print(f"Row {ri}: {row_text}")

# print("\n==========================")
# print(" 🔍 TEST 3: detect keys")
# print("==========================")
# keys = ["customer", "product", "priority"]

# # scan paragraphs
# print("Scan paragraphs for keys:")
# for p in doc.paragraphs:
#     t = p.text.strip().lower()
#     for k in keys:
#         if k in t:
#             print(f"Found '{k}' in paragraph: '{p.text}'")

# # scan tables
# print("\nScan tables for keys:")
# for table in doc.tables:
#     for row in table.rows:
#         for cell in row.cells:
#             t = cell.text.strip().lower()
#             for k in keys:
#                 if k in t:
#                     print(f"Found '{k}' in table cell: '{cell.text}'")

# print("\n==========================")
# print(" END TEST ")
# print("==========================")
