import requests
import json
from docx import Document
import os
import getpass
from dotenv import load_dotenv
from langsmith import traceable
import re

# Read key from .env
load_dotenv()
# Manual input
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if "OPENAI_API_KEY" not in os.environ:
    os.environ["OPENAI_API_KEY"] = getpass.getpass("Enter your OpenAI API key: ")

os.environ["LANGSMITH_TRACING"] = "true"
os.environ["LANGSMITH_PROJECT"] = "8D-LLM-Monitoring"

@traceable(run_type="llm", name="8D-Section-Extraction")
def call_litellm(prompt, model="azure/gpt-4.1", api_base="http://litellm.ame.local"):
    url = f"{api_base}/v1/chat/completions"

    headers = {
        "Authorization": f"Bearer {OPENAI_API_KEY}",
        "Content-Type": "application/json",
    }

    payload = {
        "model": model,
        "messages": [
            {"role": "user", "content": prompt}
        ],
        "temperature": 0,
        "stream": False,
    }

    resp = requests.post(url, json=payload, headers=headers)
    resp.raise_for_status()
    data = resp.json()
    return data["choices"][0]["message"]["content"]

def safe_json_loads(text):
    text = text.strip()

    # remove leading ```json or ``` 
    text = re.sub(r"^```json\s*", "", text)
    text = re.sub(r"^```\s*", "", text)

    # remove trailing ```
    text = re.sub(r"\s*```$", "", text)

    # trim again
    text = text.strip()

    # If still contains extra text, crop between first { and last }
    if "{" in text and "}" in text:
        text = text[text.find("{"): text.rfind("}")+1]

    return json.loads(text)

def split_by_headings(doc_path, keep_levels=("Heading 1", "Heading 2")):
    doc = Document(doc_path)
    sections = []
    current_title = None
    current_text = []

    for para in doc.paragraphs:
        if para.style.name in keep_levels:
            # Save previous section
            if current_title is not None:
                sections.append({
                    "title": current_title,
                    "content": "\n".join(current_text).strip()
                })
                current_text = []

            current_title = para.text.strip()
        else:
            if current_title is not None:
                current_text.append(para.text)

    # Save last section
    if current_title:
        sections.append({
            "title": current_title,
            "content": "\n".join(current_text).strip()
        })

    return sections

def extract_product_from_tables(doc_path):
    """
    Read all tables and extract the value from a row where the first cell starts with 'Product'.
    This works ONLY when the Introduction is a real table (not a textbox).
    """
    doc = Document(doc_path)

    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue

            key = row.cells[0].text.strip().lower()
            value = row.cells[1].text.strip()

            # Detect product-like fields
            if key.startswith("product"):
                return value

    return None  # No product found in tables


def extract_problem_from_D2(raw_text):
    prompt = f"""
You are extracting structured engineering information from an 8D report.

SECTION: D2 - Problem Symptoms / Failure Mode  
TEXT:
{raw_text}

Return ONLY a JSON object with the following field:

{{
  "problem_symptoms": "<summary strictly based on the input text>"
}}
"""
    result = call_litellm(prompt)
    return json.loads(result)

def extract_root_cause_from_D4(raw_text):
    prompt = f"""
You are extracting structured engineering information from an 8D report.

SECTION: D4 - Root Cause Analysis  
TEXT:
{raw_text}

Return ONLY a JSON object with the following field:

{{
  "root_cause": "<the single most important root cause, strictly derived from input text>"
}}
"""
    result = call_litellm(prompt)
    return safe_json_loads(result)

def build_8d_json(doc_path):
    product_name = extract_product_from_tables(doc_path)
    print(product_name)
    sections = split_by_headings(doc_path)

    result = {
        "product_name": product_name,
        "sections": {},
        "raw_context": {}
    }

    for sec in sections:
        title = sec["title"]
        content = sec["content"]

        # # INTRODUCTION → extract product
        # if "Introduction" in title:
        #     result["product_name"] = extract_product_name(content)

        # D2: Problem / Failure Mode
        if title.startswith("D2"):
            result["raw_context"]["D2"] = content
            extracted = extract_problem_from_D2(content)
            result["sections"]["D2"] = extracted

        # D4: Root Cause
        elif title.startswith("D4"):
            result["raw_context"]["D4"] = content
            extracted = extract_root_cause_from_D4(content)
            result["sections"]["D4"] = extracted

    return result



if __name__ == "__main__":
    doc_path =r"C:\Users\FW\Desktop\FMEA_AI\Project_Phase\DATA\8D\8D6264240043R03.docx"  # <--- change this
    schema = build_8d_json(doc_path)

    with open("8d_output.json", "w", encoding="utf-8") as f:
        json.dump(schema, f, indent=2, ensure_ascii=False)

    print("8D JSON schema generated: 8d_output.json")