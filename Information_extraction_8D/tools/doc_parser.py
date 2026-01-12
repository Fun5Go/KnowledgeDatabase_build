# tools/doc_parser.py

from docx import Document
import re
from Information_extraction_8D.Schemas.eightD_schema_json import EightDCase
import json
from pydantic import ValidationError

from docx import Document

def split_sections(doc_path):
    # Load the .docx document from disk
    doc = Document(doc_path)

    # This will hold the final list of sections:
    # each item is {"title": <Heading 2 text>, "content": <joined paragraph text>}
    sections = []

    # Tracks the current section title (a Heading 2)
    current_title = None

    # Accumulates paragraph text that belongs to the current section
    current_text = []

    # Iterate over every paragraph in the document, in order
    for para in doc.paragraphs:
        # Get the paragraph's style name (e.g., "Heading 1", "Heading 2", "Normal", etc.)
        style = para.style.name

        # If this paragraph is a Heading 2, it marks the start of a new section
        if style == "Heading 2":

            # If we were already collecting a previous section, finalize it and store it
            if current_title:
                sections.append({
                    "title": current_title,
                    # Join collected paragraphs with newlines, remove leading/trailing whitespace
                    "content": "\n".join(current_text).strip()
                })

            # Start a new section with this heading text as the title
            current_title = para.text.strip()

            # Reset the content accumulator for the new section
            current_text = []

        else:
            # For non-heading paragraphs:
            # - only collect text if we are inside a section (current_title exists)
            # - ignore empty/whitespace-only paragraphs
            if current_title and para.text.strip():
                current_text.append(para.text.strip())

    # After the loop, we may have an unfinished section to save
    if current_title:
        sections.append({
            "title": current_title,
            "content": "\n".join(current_text).strip()
        })

    # Return the list of extracted sections
    return sections


def extract_product(doc_path):
    doc = Document(doc_path)
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2: continue
            key = row.cells[0].text.lower()
            if "product" in key:
                return row.cells[1].text.strip()
    return None


def find_8d_id(filename):
    m = re.findall(r"8D\d+", filename)
    return m[0] if m else filename




def safe_json(resp: str):
    """Extract JSON from LLM response."""
    text = resp.strip()
    if text.startswith("```"):
        text = text.split("```")[1]
    start = text.find("{")
    end = text.rfind("}")
    return json.loads(text[start:end+1])

def validate_eight_d(json_obj):
    try:
        parsed = EightDCase.model_validate(json_obj)
        return parsed
    except ValidationError as e:
        return {"validation_error": str(e)}
